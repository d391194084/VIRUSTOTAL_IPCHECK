import urllib.request
import urllib.parse
import json
import sys
import os
import ipaddress
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.auth.exceptions import RefreshError  # 新增：用於精準捕捉 Token 過期錯誤

# ==========================================
# 核心功能模組：基礎設施驗證與 API 資料獲取
# ==========================================

def validate_ip(ip: str) -> bool:
    try:
        ipaddress.ip_address(ip)
        return True
    except ValueError:
        print(f"❌ 錯誤：無效的 IP 格式 '{ip}'，請檢查輸入。")
        return False

def ip_in_fplist(ip: str, fp_list: list) -> bool:
    pattern = r'(?<![0-9\.])' + re.escape(ip) + r'(?![0-9\.])'
    return bool(re.search(pattern, json.dumps(fp_list)))

def get_vt_data(ip):
    vt_key = os.environ.get('VT_API_KEY')
    if not vt_key:
        return "❌ 錯誤：找不到 VT_API_KEY。"
        
    base_url = f"https://www.virustotal.com/api/v3/ip_addresses/{ip}"
    req = urllib.request.Request(base_url)
    req.add_header('accept', 'application/json')
    req.add_header('x-apikey', vt_key.strip())
    
    try:
        response = urllib.request.urlopen(req)
        data = json.loads(response.read())['data']['attributes']
        
        stats = data.get('last_analysis_stats', {})
        asn = data.get('asn', 'Unknown')
        as_owner = data.get('as_owner', 'Unknown')
        tags = ", ".join(data.get('tags', []))
        
        return f"""
        狀態: 成功獲取 VT 數據
        目標 IP: {ip}
        地理位置: {data.get('country', 'Unknown')}
        VT 偵測: {stats.get('malicious', 0)} / {sum(stats.values())} (Malicious/Total)
        標籤: {tags}
        ASN 背景: {as_owner} (AS{asn})
        """
    except Exception as e:
        return f"狀態: VT 查詢失敗或無回應 ({e})"

def check_false_positive(ip):
    tf_key = os.environ.get('THREATFOX_API_KEY')
    if not tf_key:
        return "⚠️ 未設定 Abuse.ch 金鑰，跳過白名單檢查"

    url = "https://hunting-api.abuse.ch/api/v1/"
    payload = {"query": "get_fplist", "format": "json"}
    data = json.dumps(payload).encode('utf-8')

    req = urllib.request.Request(url, data=data)
    req.add_header('Content-Type', 'application/json')
    req.add_header('Auth-Key', tf_key.strip())
    req.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)')

    try:
        resp = urllib.request.urlopen(req, timeout=15)
        res = json.loads(resp.read())

        if res.get('query_status') == 'ok':
            fp_list = res.get('data', [])
            if ip_in_fplist(ip, fp_list):
                return f"✅ 【安全確認】此 IP ({ip}) 已被 Abuse.ch 官方明確列為 False Positive (誤報白名單)！請大幅降低其風險評級。"
            else:
                return "不在 Abuse.ch 官方誤報白名單中 (需依賴其他情資判斷)"
        else:
            return f"⚠️ 獲取白名單失敗: {res.get('query_status')}"
    except Exception as e:
        return f"⚠️ 白名單查詢異常 ({e})"

def get_abuse_ch_data(ip):
    tf_key = os.environ.get('THREATFOX_API_KEY')
    tf_result_text = "ℹ️ 未設定 ThreatFox 金鑰，跳過家族比對"
    urlhaus_result_text = "✅ 查無惡意載體分發紀錄"
    
    # --- 1. ThreatFox 查詢 ---
    if tf_key:
        try:
            url_tf = "https://threatfox-api.abuse.ch/api/v1/"
            payload_tf = {"query": "search_ioc", "search_term": ip}
            data_tf = json.dumps(payload_tf).encode('utf-8')
            
            req_tf = urllib.request.Request(url_tf, data=data_tf)
            req_tf.add_header('Content-Type', 'application/json')
            req_tf.add_header('Auth-Key', tf_key.strip())
            req_tf.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)')
            
            with urllib.request.urlopen(req_tf, timeout=15) as resp_tf:
                res_tf = json.loads(resp_tf.read())
            
            if res_tf.get('query_status') == 'ok':
                findings = []
                for doc in res_tf.get('data', []):
                    malware = doc.get('malware_printable', '未知')
                    confidence = doc.get('confidence_level', 0)
                    findings.append(f"● 命中家族: {malware} (置信度: {confidence}%)")
                tf_result_text = "🔍 偵測到明確威脅關聯：\n      " + "\n      ".join(set(findings))
            elif res_tf.get('query_status') == 'no_result':
                tf_result_text = "✅ ThreatFox 數據庫中查無匹配紀錄。"
            else:
                tf_result_text = f"⚠️ 查詢狀態：{res_tf.get('query_status')}"
        except Exception as e:
            tf_result_text = f"⚠️ ThreatFox 查詢中斷 ({str(e)})"

    # --- 2. URLhaus 查詢 ---
    try:
        url_uh = "https://urlhaus-api.abuse.ch/v1/host/"
        data_uh = urllib.parse.urlencode({"host": ip}).encode('utf-8')
        req_uh = urllib.request.Request(url_uh, data=data_uh)
        req_uh.add_header('User-Agent', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)')
        
        # 修正點：確保 res_uh 在 try 區塊內正確賦值，避免 NameError
        with urllib.request.urlopen(req_uh, timeout=15) as resp_uh:
            res_uh = json.loads(resp_uh.read())
        
        if res_uh.get('query_status') == 'ok':
            urlhaus_result_text = f"⚠️ 發現 {len(res_uh.get('urls', []))} 筆惡意網址關聯紀錄。"
        else:
            urlhaus_result_text = "✅ URLhaus 查無此主機之惡意活動紀錄。"
            
    except urllib.error.HTTPError as e:
        if e.code == 401:
            urlhaus_result_text = "💡 URLhaus 提示：API 金鑰無效或權限受限。"
        else:
            urlhaus_result_text = f"💡 URLhaus 查詢受限 (HTTP {e.code})"
    except Exception as e:
        urlhaus_result_text = f"⚠️ URLhaus 數據分析異常 ({str(e)})"
        
    return f"\n    [ThreatFox 情資]: {tf_result_text}\n    [URLhaus 行為紀錄]: {urlhaus_result_text}\n    "
    
def check_firehol_l3(ip: str) -> str:
    """新增：動態下載 FireHOL Level 3 清單並比對 IP"""
    url = "https://iplists.firehol.org/files/firehol_level3.netset"
    try:
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'})
        resp = urllib.request.urlopen(req, timeout=10)
        data = resp.read().decode('utf-8')
        target = ipaddress.ip_address(ip)
        
        for line in data.splitlines():
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            network = ipaddress.ip_network(line, strict=False)
            if target in network:
                return f"🚨 命中！此 IP 被列入 FireHOL Level 3 黑名單 (匹配網段: {line})"
                
        return "✅ 無命中紀錄 (不在 FireHOL Level 3 中)"
    except Exception as e:
        return f"⚠️ FireHOL 查詢異常 ({e})"

def get_abuseipdb_data(ip: str) -> str:
    """
    獲取 AbuseIPDB 濫用情資 (使用 Free API Key)
    功能：取得該 IP 的濫用信心指數、總回報次數、ISP 資訊與最近回報日期。
    """
    api_key = os.environ.get('ABUSEIPDB_API_KEY')
    if not api_key:
        return "⚠️ 未設定 ABUSEIPDB_API_KEY，跳過 AbuseIPDB 查詢。"

    # API 參數設定
    url = 'https://api.abuseipdb.com/api/v2/check'
    querystring = {
        'ipAddress': ip,
        'maxAgeInDays': '90',  # 查詢最近 90 天內的回報
        'verbose': 'true'
    }
    
    headers = {
        'Accept': 'application/json',
        'Key': api_key.strip()
    }

    try:
        # 使用標準庫 urllib 進行請求，保持與你原始程式碼一致的風格
        full_url = f"{url}?{urllib.parse.urlencode(querystring)}"
        req = urllib.request.Request(full_url, headers=headers)
        
        with urllib.request.urlopen(req, timeout=15) as response:
            res_data = json.loads(response.read())
            data = res_data.get('data', {})
            
            # 提取關鍵欄位
            score = data.get('abuseConfidenceScore', 0)
            total_reports = data.get('totalReports', 0)
            usage_type = data.get('usageType', 'Unknown')
            isp = data.get('isp', 'Unknown')
            domain = data.get('domain', 'Unknown')
            last_reported = data.get('lastReportedAt', 'N/A')

            # 根據分數給予視覺化標記
            if score >= 75:
                status_icon = "🚨 [極高風險]"
            elif score >= 25:
                status_icon = "⚠️ [可疑活動]"
            else:
                status_icon = "✅ [信譽良好]"

            return f"""
        狀態: {status_icon}
        濫用信心指數: {score}%
        最近 90 天回報總數: {total_reports} 次
        最後回報時間: {last_reported}
        ISP/組織: {isp} ({domain})
        使用類型: {usage_type}
            """
    except urllib.error.HTTPError as e:
        if e.code == 429:
            return "⚠️ AbuseIPDB 查詢失敗: 已達到每日 API 配額上限 (Rate Limit)。"
        return f"⚠️ AbuseIPDB 查詢失敗 (HTTP {e.code})"
    except Exception as e:
        return f"⚠️ AbuseIPDB 查詢異常 ({e})"

# ==========================================
# 智慧引擎與排版模組
# ==========================================

def analyze_with_gemini(combined_data):
    print("🧠 [2/4] 正在向 Google 索取可用模型總表並執行全自動闖關...")
    
    api_key = os.environ.get('GEMINI_API_KEY')
    if not api_key:
        print("❌ 錯誤：找不到 GEMINI_API_KEY。")
        sys.exit(1)
        
    api_key = api_key.strip()

    list_url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
    try:
        req_list = urllib.request.Request(list_url)
        resp_list = urllib.request.urlopen(req_list)
        models_data = json.loads(resp_list.read())
        
        available_models = [
            m['name'] for m in models_data.get('models', [])
            if 'generateContent' in m.get('supportedGenerationMethods', [])
            and 'gemini' in m.get('name', '').lower()
        ]
        print(f"   📋 系統回報：您的金鑰帳面上共有 {len(available_models)} 個潛在可用模型。")
    except Exception as e:
        print(f"❌ 獲取模型清單失敗: {e}")
        sys.exit(1)

    preferred = [
        "models/gemini-2.5-flash", 
        "models/gemini-2.0-flash", 
        "models/gemini-1.5-flash", 
        "models/gemini-1.5-pro", 
        "models/gemini-pro"
    ]
    
    prioritized_models = [m for m in preferred if m in available_models]
    for m in available_models:
        if m not in prioritized_models:
            prioritized_models.append(m)

    tw_tz = timezone(timedelta(hours=8))
    current_date = datetime.now(tw_tz).strftime('%Y年%m月%d日')

    prompt = f"""
    今日日期：{current_date}（台灣標準時間）
    你是一位專業的資安分析顧問。請根據以下多源情資，為企業資訊部門產出一份客觀、精確的分析報告。
    報告日期請使用上方提供的今日日期，不得自行推測或捏造。
    
    【重要原則】
    1. 實事求是：僅針對有證據的威脅進行分析。若數據顯示為「無命中」或「API 報錯」，請如實說明為「查無紀錄」或「技術性缺口」，不要過度推測尚未發現的威脅。
    2. 語氣專業：使用客觀、穩重的商務繁體中文，避免誇張或威脅性的措辭。
    3. 證據權重：若 AbuseIPDB 分數高但 VirusTotal 偵測為 0，請分析其可能為新興掃描活動或誤報，而非直接判定為高風險。

    【情資數據】
    {combined_data}

    【輸出格式要求】
    一、 綜合風險評估 (請給出 High/Medium/Low/Info 並說明核心理由)
    二、 情資數據摘要 (包含數據源命中狀況)
    三、 技術分析 (針對 IP 背景與行為進行中性解讀)
    四、 觀察與建議行動
    """
    
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    data = json.dumps(payload).encode('utf-8')

    for model_name in prioritized_models:
        print(f"   ⏳ 嘗試呼叫最佳模型: {model_name} ...")
        
        url = f"https://generativelanguage.googleapis.com/v1beta/{model_name}:generateContent?key={api_key}"
        
        req = urllib.request.Request(url, data=data)
        req.add_header('Content-Type', 'application/json')
        
        try:
            response = urllib.request.urlopen(req)
            result = json.loads(response.read())
            print(f"   ✅ 闖關成功！最終為您完成分析的模型是：{model_name}")
            return result['candidates'][0]['content']['parts'][0]['text']
            
        except urllib.error.HTTPError as e:
            try:
                error_info = json.loads(e.read().decode())
                err_msg = error_info.get('error', {}).get('message', '未知錯誤')
            except:
                err_msg = str(e)
            print(f"   ⚠️ 失敗 ({e.code}): {err_msg}")
            continue
        except Exception as e:
            print(f"   ⚠️ 發生未知錯誤: {e}")
            continue

    print("❌ 致命錯誤：清單內所有模型皆被 Google 伺服器拒絕存取。請確認您的 API Key 是否有效。")
    sys.exit(1)

def extract_risk_level(content: str) -> str:
    for level in ['High', 'Medium', 'Low']:
        if level.lower() in content.lower():
            return level
    return 'Unknown'

def create_word_document(ip, content):
    print("📝 [3/4] 正在生成企業級 Word (.docx) 報告...")

    # ── 顏色常數 ────────────────────────────────────────────────────────────────
    C_NAVY       = '1F3864'   # 深藍：標題列、標籤色
    C_MID_BLUE   = '2E75B6'   # 中藍：次要標題
    C_LABEL_BG   = 'D6E4F0'   # 淡藍：Label 欄背景
    C_DIVIDER    = 'BDD7EE'   # 淡藍灰：水平分隔線
    C_TEXT_DARK  = '1A1A2E'   # 深字色
    C_TEXT_LIGHT = 'FFFFFF'   # 白字
    C_RISK = {
        'High':    ('C00000', C_TEXT_LIGHT),  # 深紅
        'Medium':  ('E36C09', C_TEXT_LIGHT),  # 橘
        'Low':     ('375623', C_TEXT_LIGHT),  # 深綠
        'Unknown': ('595959', C_TEXT_LIGHT),  # 灰
    }

    # ── XML Helper：設定儲存格背景色 ─────────────────────────────────────────
    def _set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    # ── XML Helper：設定段落背景色 ───────────────────────────────────────────
    def _set_para_bg(para, hex_color):
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        pPr.append(shd)

    # ── XML Helper：加入段落左側色條 ─────────────────────────────────────────
    def _set_left_border(para, color, sz='20', space='6'):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    sz)
        left.set(qn('w:space'), space)
        left.set(qn('w:color'), color)
        pBdr.append(left)
        pPr.append(pBdr)

    # ── XML Helper：加入段落下底線（水平分隔） ───────────────────────────────
    def _set_bottom_border(para, color, sz='6'):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    sz)
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Helper：加入章節標題段落 ─────────────────────────────────────────────
    def _add_section_heading(doc, text):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after  = Pt(4)
        para.paragraph_format.left_indent  = Pt(10)
        _set_left_border(para, C_NAVY, sz='28', space='8')
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Microsoft JhengHei'
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        return para

    # ── Helper：加入一般內文段落 ─────────────────────────────────────────────
    def _add_body_para(doc, text):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Pt(18)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = 'Microsoft JhengHei'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        return para

    # ── 建立 Document ────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.2)
    sec.right_margin  = Inches(1.2)

    # ─────────────────────────────────────────────────────────────────────────
    # 區塊 1：深藍標題橫幅
    # ─────────────────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(0)
    _set_para_bg(p_title, C_NAVY)
    r = p_title.add_run('  資安威脅深度分析報告  ')
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = 'Microsoft JhengHei'
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(14)
    _set_para_bg(p_sub, C_NAVY)
    r2 = p_sub.add_run('Cybersecurity Threat Intelligence Report')
    r2.font.size = Pt(11)
    r2.font.name = 'Arial'
    r2.font.color.rgb = RGBColor(0xBF, 0xD7, 0xED)

    # ─────────────────────────────────────────────────────────────────────────
    # 區塊 2：元資料彙整表格
    # ─────────────────────────────────────────────────────────────────────────
    tw_tz      = timezone(timedelta(hours=8))
    risk_level = extract_risk_level(content)
    risk_bg, risk_fg = C_RISK.get(risk_level, C_RISK['Unknown'])

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'

    # ── 表格欄寬（DXA，約 8.76 cm + 9 cm） ─────────────────────────────────
    # 頁面可用寬度 ≈ 8.5" – 2.4" = 6.1" = 8784 DXA；各欄 3888 / 4896
    col_widths = [3888, 4896]
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(col_widths[j]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

    # ── Row 0：標頭列（合併儲存格） ──────────────────────────────────────────
    hdr = tbl.rows[0]
    hdr.cells[0].merge(hdr.cells[1])
    _set_cell_bg(hdr.cells[0], C_NAVY)
    hp = hdr.cells[0].paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(5)
    hp.paragraph_format.space_after  = Pt(5)
    hr_ = hp.add_run('▌ 報告基本資訊  /  Report Metadata')
    hr_.bold = True
    hr_.font.size = Pt(11.5)
    hr_.font.name = 'Microsoft JhengHei'
    hr_.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Rows 1-3：資料列 ─────────────────────────────────────────────────────
    meta_rows = [
        ('評估目標  Target IP',     ip),
        ('報告產出時間  Generated', datetime.now(tw_tz).strftime('%Y-%m-%d %H:%M:%S') + '（台灣標準時間 UTC+8）'),
        ('綜合風險等級  Risk Level', f'  ●  {risk_level.upper()}  '),
    ]
    for i, (label, value) in enumerate(meta_rows):
        row   = tbl.rows[i + 1]
        lcell = row.cells[0]
        vcell = row.cells[1]

        # Label 欄
        _set_cell_bg(lcell, C_LABEL_BG)
        lp = lcell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(5)
        lp.paragraph_format.space_after  = Pt(5)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.name = 'Microsoft JhengHei'
        lr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # Value 欄
        vp = vcell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(5)
        vp.paragraph_format.space_after  = Pt(5)
        if i == 2:  # 風險等級：加底色
            _set_cell_bg(vcell, risk_bg)
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = vp.add_run(value)
            vr.bold = True
            vr.font.size = Pt(11)
            vr.font.name = 'Microsoft JhengHei'
            fg_r, fg_g, fg_b = (int(risk_fg[j:j+2], 16) for j in (0, 2, 4))
            vr.font.color.rgb = RGBColor(fg_r, fg_g, fg_b)
        else:
            _set_cell_bg(vcell, 'FFFFFF')
            vr = vp.add_run(value)
            vr.font.size = Pt(10)
            vr.font.name = 'Microsoft JhengHei'
            vr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()   # 表格後留白

    # ─────────────────────────────────────────────────────────────────────────
    # 區塊 3：AI 分析正文
    # ─────────────────────────────────────────────────────────────────────────
    section_markers = ('執行摘要', '一、', '二、', '三、', '四、', '五、', '六、')
    bullet_prefixes = ('●', '·', '-', '*', '•', '○', '▶', '→')

    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue

        is_heading = False
        for m in section_markers:
            if stripped.startswith(m):
                _add_section_heading(doc, stripped)
                is_heading = True
                break

        if not is_heading:
            if any(stripped.startswith(pfx) for pfx in bullet_prefixes):
                body = stripped[1:].strip()
                para = doc.add_paragraph()
                para.paragraph_format.left_indent  = Pt(28)
                para.paragraph_format.first_line_indent = Pt(-14)
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                bullet_run = para.add_run('▸  ')
                bullet_run.bold = True
                bullet_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                bullet_run.font.size = Pt(10.5)
                text_run = para.add_run(body)
                text_run.font.size = Pt(10.5)
                text_run.font.name = 'Microsoft JhengHei'
                text_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            else:
                _add_body_para(doc, stripped)

    # ─────────────────────────────────────────────────────────────────────────
    # 區塊 4：頁腳資訊列
    # ─────────────────────────────────────────────────────────────────────────
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(18)
    p_div.paragraph_format.space_after  = Pt(4)
    _set_bottom_border(p_div, C_NAVY, sz='12')

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(2)
    p_foot.paragraph_format.space_after  = Pt(0)
    fr = p_foot.add_run(
        f'本報告屬機密文件，僅供授權人員閱覽  ·  '
        f'Auto Analyst v1.0  ·  '
        f'Generated: {datetime.now(tw_tz).strftime("%Y-%m-%d")}'
    )
    fr.italic = True
    fr.font.size = Pt(8.5)
    fr.font.name = 'Arial'
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = f"Security_Report_{ip.replace('.', '_')}.docx"
    doc.save(filename)
    return filename

def upload_to_drive(filename):
    print("☁️ [4/4] 正在使用您本人的專屬授權將報告上傳至 Google Drive...")
    
    client_id = os.environ.get('GDRIVE_CLIENT_ID')
    client_secret = os.environ.get('GDRIVE_CLIENT_SECRET')
    refresh_token = os.environ.get('GDRIVE_REFRESH_TOKEN')
    folder_id = os.environ.get('GDRIVE_FOLDER_ID')
    
    if not all([client_id, client_secret, refresh_token, folder_id]):
        print("❌ 錯誤：缺少 Google Drive OAuth 相關的環境變數！")
        sys.exit(1)

    creds = Credentials(
        token=None,
        refresh_token=refresh_token.strip(),
        token_uri="https://oauth2.googleapis.com/token",
        client_id=client_id.strip(),
        client_secret=client_secret.strip()
    )
    
    service = build('drive', 'v3', credentials=creds)
    
    file_metadata = {'name': filename, 'parents': [folder_id.strip()]}
    media = MediaFileUpload(filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    try:
        file = service.files().create(
            body=file_metadata, media_body=media, fields='id', supportsAllDrives=True
        ).execute()
        print(f"✅ 完美登頂！報告已成功存入您的 Google Drive，檔案 ID: {file.get('id')}")
    except RefreshError as e:
        print(f"❌ 錯誤：Google Drive 授權已失效 ({e})。請重新發行 Refresh Token 並更新至環境變數中。")
        sys.exit(1)

# ==========================================
# 主程式執行區塊
# ==========================================

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("用法: python auto_analyst.py <IP地址>")
        sys.exit(1)
        
    target_ip = sys.argv[1]
    
    if not validate_ip(target_ip):
        sys.exit(1)
    
    # 修改原本的併行執行部分
    print("⚡ 🔍 [1/4] 啟動 5X 引擎：正在並行獲取 VT、Abuse.ch、FireHOL 與 AbuseIPDB 情資...")
    with ThreadPoolExecutor(max_workers=5) as ex:
        f_vt      = ex.submit(get_vt_data, target_ip)
        f_fp      = ex.submit(check_false_positive, target_ip)
        f_abuse   = ex.submit(get_abuse_ch_data, target_ip)
        f_firehol = ex.submit(check_firehol_l3, target_ip)
        f_abuseipdb = ex.submit(get_abuseipdb_data, target_ip) # <--- 新增這行
        
    # 獲取所有結果
    vt_info      = f_vt.result()
    fp_info      = f_fp.result()
    abuse_info   = f_abuse.result()
    firehol_info = f_firehol.result()
    abuseipdb_info = f_abuseipdb.result() # <--- 新增這行
    
    # 將數據餵入 combined_intel 變數
    combined_intel = f"""
    --- VirusTotal 數據 ---
    {vt_info}
    
    --- AbuseIPDB 社群回報數據 ---
    {abuseipdb_info}
    
    --- Abuse.ch 誤報白名單 (False Positive) 檢查 ---
    狀態: {fp_info}
    
    --- Abuse.ch (ThreatFox + URLhaus) 惡意數據 ---
    {abuse_info}

    --- FireHOL Level 3 黑名單比對 ---
    狀態: {firehol_info}
    """
    
    report_text = analyze_with_gemini(combined_intel)
    doc_name = create_word_document(target_ip, report_text)
    upload_to_drive(doc_name)
